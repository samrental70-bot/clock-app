from __future__ import annotations

import math
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT_DIR = Path(r"C:\Users\samra\Dropbox\CompanyCam Projects")
OUTPUT_DIR = ROOT_DIR / "Portfolio Documents"
ASSET_DIR = OUTPUT_DIR / "_final_only_assets"
OUTPUT_DOCX = OUTPUT_DIR / "final-projects-portfolio-final-only.docx"

COMPANY_NAME = "Ottawa Renovation Pro Ltd."
PORTFOLIO_TITLE = "Completed Projects Portfolio"
PORTFOLIO_SUBTITLE = "Final renovation photography organized by bathrooms, kitchens, basements, and flooring."
ISSUE_DATE = "July 2026"

PAGE_SIZE = (1700, 2200)
PAGE_MARGIN = 110
PHOTO_RADIUS = 18

PAPER = "#F6F1EA"
INK = "#1B2D33"
MUTED = "#687177"
ACCENT = "#A96D34"
LINE = "#D8CEC1"
DARK_PANEL = "#F3ECE2"
WHITE = "#FFFFFF"
SOFT_FILL = "#EFE5D8"

FONT_SANS = Path(r"C:\Windows\Fonts\segoeui.ttf")
FONT_SANS_BOLD = Path(r"C:\Windows\Fonts\segoeuib.ttf")
FONT_DISPLAY = Path(r"C:\Windows\Fonts\bahnschrift.ttf")


@dataclass(frozen=True)
class PortfolioImage:
    project: str
    caption: str
    path: Path
    crop: tuple[float, float, float, float] = (0.0, 0.0, 1.0, 1.0)


def load_font(path: Path, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    try:
        return ImageFont.truetype(str(path), size=size)
    except OSError:
        return ImageFont.load_default()


DISPLAY_60 = load_font(FONT_DISPLAY, 60)
DISPLAY_44 = load_font(FONT_DISPLAY, 44)
DISPLAY_30 = load_font(FONT_DISPLAY, 30)
SANS_28 = load_font(FONT_SANS, 28)
SANS_24 = load_font(FONT_SANS, 24)
SANS_22 = load_font(FONT_SANS, 22)
SANS_20 = load_font(FONT_SANS, 20)
SANS_18 = load_font(FONT_SANS, 18)
SANS_BOLD_22 = load_font(FONT_SANS_BOLD, 22)
SANS_BOLD_18 = load_font(FONT_SANS_BOLD, 18)
SANS_BOLD_16 = load_font(FONT_SANS_BOLD, 16)


def pi(
    folder: str,
    filename: str,
    project: str,
    caption: str,
    crop: tuple[float, float, float, float] = (0.0, 0.0, 1.0, 1.0),
) -> PortfolioImage:
    return PortfolioImage(
        project=project,
        caption=caption,
        path=ROOT_DIR / folder / filename,
        crop=crop,
    )


ALI_COVER = pi(
    "Ali - Orleans Ottawa ON",
    "project_96924648_01092026_1728 (1).jpg",
    "Ali | Orleans",
    "Freestanding tub and frameless glass shower",
)
ALI_WIDE = pi(
    "Ali - Orleans Ottawa ON",
    "project_96924648_01092026_1729 (1).jpg",
    "Ali | Orleans",
    "Bright primary bathroom with open floor plan",
)
ALI_TUB_SHOWER = pi(
    "Ali - Orleans Ottawa ON",
    "project_96924648_01092026_1729 (10).jpg",
    "Ali | Orleans",
    "Tub-and-shower pairing with matte black fixtures",
)
ALI_DOUBLE_VANITY = pi(
    "Ali - Orleans Ottawa ON",
    "project_96924648_01092026_1729 (2).jpg",
    "Ali | Orleans",
    "Double vanity and full-height tile floor",
    crop=(0.10, 0.02, 0.98, 0.92),
)
ALI_MUDROOM = pi(
    "Ali - Orleans Ottawa ON",
    "project_96924648_01092026_1731 (1) (1).jpg",
    "Ali | Orleans",
    "Mudroom bench and utility entry finish",
)

BATHREN_TUB = pi(
    "Bathroom Renovation",
    "project_99424318_01062024_0913.jpg",
    "Bathroom Renovation",
    "Marble-look tile surround and tub detail",
)
BATHREN_COMPACT = pi(
    "Bathroom Renovation",
    "project_99424318_02222026_1359 (12).jpg",
    "Bathroom Renovation",
    "Compact bath with custom vanity",
)
BATHREN_VANITY = pi(
    "Bathroom Renovation",
    "project_99424318_02222026_1358 (4).jpg",
    "Bathroom Renovation",
    "Double vanity with arched mirrors",
    crop=(0.08, 0.0, 1.0, 0.90),
)

KRISHNA_BATH = pi(
    "Bathroom Krishna",
    "project_99424299_10132022_0805 (2).jpg",
    "Krishna Bathroom",
    "Clean wall tile and compact bath layout",
)
KRISHNA_LOW = pi(
    "Bathroom Krishna",
    "project_99424299_02222026_1355 (4).jpg",
    "Krishna Bathroom",
    "Low-angle view emphasizing flooring and fixture finish",
)
KRISHNA_VANITY = pi(
    "Bathroom Krishna",
    "project_99424299_02222026_1355 (6).jpg",
    "Krishna Bathroom",
    "Vanity and tub wall finish",
)

ASIYA_KITCHEN = pi(
    "Basement Kitchen Bathroom Asiya new",
    "project_96110391_02222026_0849 (2).jpg",
    "Asiya Project",
    "Modern white kitchen with stone counters",
)
ASIYA_FIREPLACE = pi(
    "Basement Kitchen Bathroom Asiya new",
    "project_96110391_02222026_0848 (1).jpg",
    "Asiya Project",
    "Feature wall and electric fireplace",
)
ASIYA_LIVING = pi(
    "Basement Kitchen Bathroom Asiya new",
    "project_96110391_02222026_0848 (3).jpg",
    "Asiya Project",
    "Open lower-level living room",
)
ASIYA_BEDROOM = pi(
    "Basement Kitchen Bathroom Asiya new",
    "project_96110391_02222026_0849 (10).jpg",
    "Asiya Project",
    "Finished bedroom with workspace",
)
ASIYA_BATHROOM = pi(
    "Basement Kitchen Bathroom Asiya new",
    "project_96110391_02222026_0849.jpg",
    "Asiya Project",
    "Freestanding tub with walk-in shower",
)
ASIYA_BATHROOM_VANITY = pi(
    "Basement Kitchen Bathroom Asiya new",
    "project_96110391_02222026_0849 (12).jpg",
    "Asiya Project",
    "Bathroom vanity with full-width mirror",
)

ARNPRIOR_BATH = pi(
    "Bathroom flooring kitchen Armprior",
    "03 Arnprior 1 after.JPG",
    "Arnprior Project",
    "Bathroom tile and tub surround finish",
)
ARNPRIOR_KITCHEN = pi(
    "Bathroom flooring kitchen Armprior",
    "03 arnprior 4 after.JPG",
    "Arnprior Project",
    "Kitchen refresh with large-format tile floor",
)
ARNPRIOR_PORCH = pi(
    "Bathroom flooring kitchen Armprior",
    "03 Arnprior 2 after.jpg",
    "Arnprior Project",
    "Exterior landing tile detail",
)

ARBI_ROOM_A = pi(
    "16 Arbi basement room",
    "project_99410262_02212026_1734.jpg",
    "Arbi Basement",
    "Finished room with integrated cabinet niche",
)
ARBI_ROOM_B = pi(
    "16 Arbi basement room",
    "project_99410262_02212026_1810.jpg",
    "Arbi Basement",
    "Lower-level room with clean trim and sliding storage",
)

MISTY_STAIRS_UP = pi(
    "Flooring Misty f",
    "project_99472041_02232026_1821 (11).jpg",
    "Misty Flooring",
    "Wood stair refinish",
    crop=(0.0, 0.10, 1.0, 1.0),
)
MISTY_STAIRS_LANDING = pi(
    "Flooring Misty f",
    "project_99472041_02232026_1821 (12).jpg",
    "Misty Flooring",
    "Landing detail and nosing finish",
)
MISTY_STAIRS_DOWN = pi(
    "Flooring Misty f",
    "project_99472041_02232026_1820 (16).jpg",
    "Misty Flooring",
    "Lower stair run with matching handrail tone",
)

INDIAN_EMBASSY_FLOOR = pi(
    "Flooring Indian Embassy",
    "project_99470490_02232026_1806 (1).jpg",
    "Indian Embassy Flooring",
    "Commercial flooring installation",
    crop=(0.0, 0.05, 1.0, 0.92),
)

CARLETON_NEW_BATH_A = pi(
    "Basement Carleton New",
    "project_99423261_02222026_0827 (16).jpg",
    "Lower-Level Ensuite",
    "Frameless shower with bold tile floor",
)
CARLETON_NEW_BATH_B = pi(
    "Basement Carleton New",
    "project_99423261_02222026_0827 (23).jpg",
    "Lower-Level Ensuite",
    "Vanity, mirror, and shower aligned as a finished bath suite",
)
CARLETON_NEW_ROOM = pi(
    "Basement Carleton New",
    "project_99423261_02222026_0827 (17).jpg",
    "Finished Bedroom",
    "Bright finished bedroom with modern lower-level flooring",
)
CARLETON_NEW_KITCHENETTE_WIDE = pi(
    "Basement Carleton New",
    "project_99423261_02222026_0827 (20).jpg",
    "Suite Living Area",
    "Lower-level kitchenette framed as a clean living space",
    crop=(0.0, 0.0, 0.89, 0.91),
)
CARLETON_NEW_KITCHENETTE_DETAIL = pi(
    "Basement Carleton New",
    "project_99423261_02222026_0827 (21).jpg",
    "Compact Kitchenette",
    "Finished kitchenette cabinet wall",
    crop=(0.0, 0.0, 0.92, 1.0),
)

ELEVEN_CARLETON_OPEN = pi(
    "11 Carleton - Carleton Place Ontario",
    "project_99404987_02212026_1515 (1).jpg",
    "11 Carleton",
    "Open lower-level suite view with kitchenette and new flooring",
)
ELEVEN_CARLETON_HALL = pi(
    "11 Carleton - Carleton Place Ontario",
    "project_99404987_02212026_1515 (41).jpg",
    "11 Carleton",
    "Long suite corridor leading into the kitchenette zone",
)
CARLETON_MORE_BEDROOM = pi(
    "11 Carleton More",
    "project_99482023_02242026_0537 (1).jpg",
    "11 Carleton More",
    "Finished bedroom with clean neutral walls and plank flooring",
)
CARLETON_MORE_HALL = pi(
    "11 Carleton More",
    "project_99482023_02242026_0537 (5).jpg",
    "Suite Hallway",
    "Hallway leading into the kitchenette zone",
    crop=(0.0, 0.0, 1.0, 0.82),
)
CARLETON_MORE_KITCHENETTE_DETAIL = pi(
    "11 Carleton More",
    "project_99482023_02242026_0537 (6).jpg",
    "Compact Kitchenette",
    "Compact kitchenette wall with durable flooring and simple storage",
)
CARLETON_MORE_KITCHENETTE_WIDE = pi(
    "11 Carleton More",
    "project_99482023_02242026_0537 (7).jpg",
    "Open Suite Layout",
    "Open suite with integrated kitchenette",
    crop=(0.0, 0.0, 0.96, 0.94),
)
CARLETON_MORE_KITCHEN_WALL = pi(
    "11 Carleton More",
    "project_99482023_02242026_0538 (1).jpg",
    "Kitchenette Wall",
    "Full kitchenette wall with modern cabinetry and stainless appliances",
    crop=(0.0, 0.0, 0.86, 1.0),
)

ALL_IMAGES = (
    ALI_COVER,
    ALI_WIDE,
    ALI_TUB_SHOWER,
    ALI_DOUBLE_VANITY,
    ALI_MUDROOM,
    BATHREN_TUB,
    BATHREN_COMPACT,
    BATHREN_VANITY,
    KRISHNA_BATH,
    KRISHNA_LOW,
    KRISHNA_VANITY,
    ASIYA_KITCHEN,
    ASIYA_FIREPLACE,
    ASIYA_LIVING,
    ASIYA_BEDROOM,
    ASIYA_BATHROOM,
    ASIYA_BATHROOM_VANITY,
    ARNPRIOR_BATH,
    ARNPRIOR_KITCHEN,
    ARNPRIOR_PORCH,
    ARBI_ROOM_A,
    ARBI_ROOM_B,
    MISTY_STAIRS_UP,
    MISTY_STAIRS_LANDING,
    MISTY_STAIRS_DOWN,
    INDIAN_EMBASSY_FLOOR,
    CARLETON_NEW_BATH_A,
    CARLETON_NEW_BATH_B,
    CARLETON_NEW_ROOM,
    CARLETON_NEW_KITCHENETTE_WIDE,
    CARLETON_NEW_KITCHENETTE_DETAIL,
    ELEVEN_CARLETON_OPEN,
    ELEVEN_CARLETON_HALL,
    CARLETON_MORE_BEDROOM,
    CARLETON_MORE_HALL,
    CARLETON_MORE_KITCHENETTE_DETAIL,
    CARLETON_MORE_KITCHENETTE_WIDE,
    CARLETON_MORE_KITCHEN_WALL,
)


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def verify_inputs() -> None:
    missing = [image.path for image in ALL_IMAGES if not image.path.exists()]
    if missing:
        joined = "\n".join(str(path) for path in missing[:12])
        raise SystemExit(f"Missing image files:\n{joined}")


def _set_rfonts(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def set_run_style(
    run,
    *,
    font_name: str,
    size: int,
    color: str,
    bold: bool = False,
    italic: bool = False,
) -> None:
    _set_rfonts(run, font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 0, line: float = 1.0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number_field(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    _set_rfonts(run, "Segoe UI")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(placeholder)
    r_element.append(fld_end)


def set_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)

    label_run = paragraph.add_run(f"{COMPANY_NAME} | Page ")
    set_run_style(label_run, font_name="Segoe UI", size=9, color=MUTED)
    add_page_number_field(paragraph)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("232B30")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[index : index + 2], 16) for index in (0, 2, 4))


def new_canvas() -> tuple[Image.Image, ImageDraw.ImageDraw]:
    canvas = Image.new("RGB", PAGE_SIZE, PAPER)
    return canvas, ImageDraw.Draw(canvas)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    if not words:
        return []

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        trial = f"{current} {word}"
        if draw.textlength(trial, font=font) <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_paragraph(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.ImageFont,
    fill: str,
    x: int,
    y: int,
    max_width: int,
    line_gap: int,
) -> int:
    lines = wrap_text(draw, text, font, max_width)
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        bbox = draw.textbbox((x, y), line, font=font)
        y += (bbox[3] - bbox[1]) + line_gap
    return y


def rounded_mask(size: tuple[int, int], radius: int) -> Image.Image:
    mask = Image.new("L", size, 0)
    mask_draw = ImageDraw.Draw(mask)
    mask_draw.rounded_rectangle((0, 0, size[0], size[1]), radius=radius, fill=255)
    return mask


def prepare_photo(image: PortfolioImage, size: tuple[int, int]) -> Image.Image:
    with Image.open(image.path) as src:
        base = ImageOps.exif_transpose(src).convert("RGB")

    left = max(0, min(base.width, int(base.width * image.crop[0])))
    top = max(0, min(base.height, int(base.height * image.crop[1])))
    right = max(left + 1, min(base.width, int(base.width * image.crop[2])))
    bottom = max(top + 1, min(base.height, int(base.height * image.crop[3])))
    cropped = base.crop((left, top, right, bottom))
    return ImageOps.fit(cropped, size, method=Image.Resampling.LANCZOS, centering=(0.5, 0.5))


def paste_photo(canvas: Image.Image, draw: ImageDraw.ImageDraw, image: PortfolioImage, box: tuple[int, int, int, int]) -> None:
    left, top, right, bottom = box
    width = right - left
    height = bottom - top

    fitted = prepare_photo(image, (width, height))
    mask = rounded_mask((width, height), PHOTO_RADIUS)
    canvas.paste(fitted, (left, top), mask)
    draw.rounded_rectangle(box, radius=PHOTO_RADIUS, outline=LINE, width=2)


def add_photo_label(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], image: PortfolioImage) -> None:
    left, _, right, bottom = box
    band_height = 92
    label_box = (left + 18, bottom - band_height - 18, right - 18, bottom - 18)
    draw.rounded_rectangle(label_box, radius=16, fill=DARK_PANEL)
    draw.line((label_box[0] + 18, label_box[1] + 14, label_box[0] + 120, label_box[1] + 14), fill=ACCENT, width=3)
    draw.text((label_box[0] + 18, label_box[1] + 22), image.project.upper(), font=SANS_BOLD_16, fill=INK)
    draw.text((label_box[0] + 18, label_box[1] + 48), image.caption, font=SANS_18, fill=MUTED)


def draw_kicker(draw: ImageDraw.ImageDraw, text: str, x: int, y: int) -> None:
    draw.text((x, y), text.upper(), font=SANS_BOLD_16, fill=ACCENT)


def draw_top_rule(draw: ImageDraw.ImageDraw) -> None:
    draw.line((PAGE_MARGIN, 70, PAGE_SIZE[0] - PAGE_MARGIN, 70), fill=LINE, width=2)


def add_cover_page() -> Path:
    canvas, draw = new_canvas()
    draw_top_rule(draw)

    hero_box = (90, 110, PAGE_SIZE[0] - 90, 1440)
    paste_photo(canvas, draw, ALI_COVER, hero_box)

    panel_box = (160, 1260, 1160, 1835)
    draw.rounded_rectangle(panel_box, radius=26, fill=PAPER, outline=LINE, width=2)
    draw_kicker(draw, COMPANY_NAME, panel_box[0] + 48, panel_box[1] + 44)

    title_y = draw_paragraph(
        draw,
        PORTFOLIO_TITLE,
        DISPLAY_60,
        INK,
        panel_box[0] + 48,
        panel_box[1] + 105,
        800,
        12,
    )
    subtitle_y = draw_paragraph(
        draw,
        PORTFOLIO_SUBTITLE,
        SANS_24,
        MUTED,
        panel_box[0] + 48,
        title_y + 16,
        800,
        10,
    )
    draw.text((panel_box[0] + 48, subtitle_y + 34), ISSUE_DATE, font=SANS_BOLD_18, fill=ACCENT)

    stat_box = (1220, 1505, 1565, 1835)
    draw.rounded_rectangle(stat_box, radius=22, fill=SOFT_FILL)
    stats = [
        ("10", "curated pages"),
        ("5", "renovation sections"),
        ("30+", "final project photos"),
    ]
    stat_y = stat_box[1] + 42
    for value, label in stats:
        draw.text((stat_box[0] + 34, stat_y), value, font=DISPLAY_30, fill=INK)
        draw.text((stat_box[0] + 120, stat_y + 8), label, font=SANS_20, fill=MUTED)
        stat_y += 82

    page_path = ASSET_DIR / "page-01-cover.png"
    canvas.save(page_path, format="PNG")
    return page_path


def add_intro_page() -> Path:
    canvas, draw = new_canvas()
    draw_top_rule(draw)
    draw_kicker(draw, "Portfolio Overview", PAGE_MARGIN, 120)

    title_y = draw_paragraph(
        draw,
        "Finished renovation work presented with a cleaner, client-ready point of view.",
        DISPLAY_44,
        INK,
        PAGE_MARGIN,
        175,
        700,
        10,
    )
    body_y = draw_paragraph(
        draw,
        "This portfolio focuses on completed spaces only. It avoids rough construction shots so the work reads the way future clients want to see it: calm, complete, and ready to live in.",
        SANS_24,
        MUTED,
        PAGE_MARGIN,
        title_y + 22,
        660,
        10,
    )
    body_y = draw_paragraph(
        draw,
        "The document is organized by bathrooms, kitchens, basements, flooring, and final detail work, with stronger coverage from Carleton and other completed lower-level suite projects.",
        SANS_24,
        MUTED,
        PAGE_MARGIN,
        body_y + 18,
        660,
        10,
    )

    category_box = (PAGE_MARGIN, body_y + 38, 700, body_y + 240)
    draw.rounded_rectangle(category_box, radius=20, fill=SOFT_FILL)
    categories = ["Bathrooms", "Kitchens", "Basements & Suites", "Flooring & Stairs", "Final Details"]
    cat_y = category_box[1] + 28
    for category in categories:
        draw.text((category_box[0] + 26, cat_y), category, font=SANS_BOLD_22, fill=INK)
        cat_y += 42

    image_box = (900, 140, 1560, 1520)
    paste_photo(canvas, draw, KRISHNA_LOW, image_box)
    add_photo_label(draw, image_box, KRISHNA_LOW)

    quote_box = (900, 1590, 1560, 1890)
    draw.rounded_rectangle(quote_box, radius=20, fill=WHITE, outline=LINE, width=2)
    quote = "Selected pages now show a wider range of finished work across bathrooms, suites, kitchenettes, lower-level rooms, stairs, and practical detail spaces."
    quote_y = draw_paragraph(draw, quote, SANS_22, INK, quote_box[0] + 28, quote_box[1] + 36, 600, 10)
    draw.text((quote_box[0] + 28, quote_y + 18), "Modern presentation focus", font=SANS_BOLD_18, fill=ACCENT)

    page_path = ASSET_DIR / "page-02-intro.png"
    canvas.save(page_path, format="PNG")
    return page_path


def add_editorial_page(
    *,
    filename: str,
    section: str,
    title: str,
    body: str,
    hero: PortfolioImage,
    supports: Iterable[PortfolioImage],
    hero_on_left: bool,
) -> Path:
    support_list = list(supports)
    if len(support_list) != 3:
        raise ValueError("Editorial pages require exactly three supporting images.")

    canvas, draw = new_canvas()
    draw_top_rule(draw)
    draw_kicker(draw, section, PAGE_MARGIN, 118)

    title_y = draw_paragraph(draw, title, DISPLAY_44, INK, PAGE_MARGIN, 172, 760, 10)
    draw_paragraph(draw, body, SANS_24, MUTED, PAGE_MARGIN, title_y + 18, 760, 10)

    hero_box = (120, 430, 1020, 1565) if hero_on_left else (680, 430, 1580, 1565)
    support_x = 1085 if hero_on_left else 120
    support_boxes = [
        (support_x, 430, support_x + 475, 760),
        (support_x, 815, support_x + 475, 1145),
        (support_x, 1200, support_x + 475, 1565),
    ]

    paste_photo(canvas, draw, hero, hero_box)
    add_photo_label(draw, hero_box, hero)

    for image, box in zip(support_list, support_boxes, strict=True):
        paste_photo(canvas, draw, image, box)
        add_photo_label(draw, box, image)

    page_path = ASSET_DIR / filename
    canvas.save(page_path, format="PNG")
    return page_path


def add_mosaic_page() -> Path:
    canvas, draw = new_canvas()
    draw_top_rule(draw)
    draw_kicker(draw, "Selected Highlights", PAGE_MARGIN, 118)

    title_y = draw_paragraph(
        draw,
        "A broader sample of finished rooms, utility spaces, and detail work across recent projects.",
        DISPLAY_44,
        INK,
        PAGE_MARGIN,
        172,
        980,
        10,
    )
    draw_paragraph(
        draw,
        "This final spread shows the range of finished work while keeping the presentation clean, calm, and easy to browse.",
        SANS_24,
        MUTED,
        PAGE_MARGIN,
        title_y + 18,
        980,
        10,
    )

    hero_box = (120, 420, 1580, 1075)
    paste_photo(canvas, draw, ASIYA_FIREPLACE, hero_box)
    add_photo_label(draw, hero_box, ASIYA_FIREPLACE)

    tiles = [
        (ALI_MUDROOM, (120, 1160, 455, 1600)),
        (ARBI_ROOM_A, (495, 1160, 830, 1600)),
        (ARNPRIOR_KITCHEN, (870, 1160, 1205, 1600)),
        (MISTY_STAIRS_LANDING, (1245, 1160, 1580, 1600)),
    ]
    for image, box in tiles:
        paste_photo(canvas, draw, image, box)
        add_photo_label(draw, box, image)

    note_box = (120, 1690, 1580, 1895)
    draw.rounded_rectangle(note_box, radius=18, fill=SOFT_FILL)
    note = (
        "Service focus: bathrooms, kitchens, basement finishing, stair and flooring upgrades, tile work, and final detail carpentry across Ottawa-area projects."
    )
    draw_paragraph(draw, note, SANS_22, INK, note_box[0] + 26, note_box[1] + 34, 1380, 8)

    page_path = ASSET_DIR / "page-09-highlights.png"
    canvas.save(page_path, format="PNG")
    return page_path


def add_closing_page() -> Path:
    canvas, draw = new_canvas()
    draw_top_rule(draw)
    summary_box = (120, 130, 1580, 720)
    draw.rounded_rectangle(summary_box, radius=28, fill=WHITE, outline=LINE, width=2)
    draw_kicker(draw, COMPANY_NAME, summary_box[0] + 36, summary_box[1] + 38)
    title_y = draw_paragraph(
        draw,
        "Completed work presented with broader project coverage, cleaner curation, and a more modern company portfolio style.",
        DISPLAY_44,
        INK,
        summary_box[0] + 36,
        summary_box[1] + 92,
        1180,
        10,
    )
    draw_paragraph(
        draw,
        "Bathrooms, kitchens, basements, lower-level suites, flooring, and final finish work across Ottawa-area residential and commercial projects. Final presentation was prioritized over rough-stage imagery so the portfolio stays client-facing and polished.",
        SANS_24,
        MUTED,
        summary_box[0] + 36,
        title_y + 18,
        1180,
        10,
    )

    coverage_box = (120, 820, 760, 1320)
    draw.rounded_rectangle(coverage_box, radius=22, fill=SOFT_FILL)
    draw.text((coverage_box[0] + 28, coverage_box[1] + 28), "Featured project mix", font=SANS_BOLD_22, fill=INK)
    coverage_lines = [
        "Ali | Orleans",
        "Asiya Basement Suite",
        "Lower-Level Ensuite",
        "Compact Kitchenette Suite",
        "Arnprior Kitchen & Bath",
        "Arbi Basement Rooms",
        "Bathroom Renovation",
        "Krishna Bathroom",
        "Misty Stair Refinishing",
        "Commercial Flooring",
    ]
    coverage_y = coverage_box[1] + 74
    for line in coverage_lines:
        draw.text((coverage_box[0] + 28, coverage_y), line, font=SANS_20, fill=INK)
        coverage_y += 36

    contact_box = (820, 820, 1580, 1180)
    draw.rounded_rectangle(contact_box, radius=22, fill=WHITE, outline=LINE, width=2)
    contact_y = contact_box[1] + 40
    for line in [
        "Phone  +1 (613) 413-1756",
        "Web    www.ottawa-renovation.ca",
        "Specialties  Bathrooms, kitchens, basements, flooring",
    ]:
        draw.text((contact_box[0] + 34, contact_y), line, font=SANS_BOLD_22, fill=INK)
        contact_y += 66

    closer_box = (820, 1245, 1580, 1575)
    draw.rounded_rectangle(closer_box, radius=22, fill=SOFT_FILL)
    closer = "Service area: Ottawa and surrounding communities. Ideal next step: send photos or a scope list to start planning your renovation."
    draw_paragraph(draw, closer, SANS_22, INK, closer_box[0] + 28, closer_box[1] + 36, 550, 8)

    note_box = (120, 1400, 760, 1765)
    draw.rounded_rectangle(note_box, radius=22, fill=WHITE, outline=LINE, width=2)
    note = "Portfolio note: selected pages emphasize final photography, broader project variety, and a cleaner client-facing presentation across the main renovation categories."
    draw_paragraph(draw, note, SANS_20, MUTED, note_box[0] + 28, note_box[1] + 34, 580, 8)

    page_path = ASSET_DIR / "page-10-closing.png"
    canvas.save(page_path, format="PNG")
    return page_path


def build_page_assets() -> list[Path]:
    return [
        add_cover_page(),
        add_intro_page(),
        add_editorial_page(
            filename="page-03-bathrooms-a.png",
            section="Bathrooms",
            title="Primary bathrooms should read like finished rooms, not isolated fixture snapshots.",
            body="This opening bathroom spread keeps the presentation bright and polished while showing tile, glazing, vanity layout, and circulation together.",
            hero=ALI_WIDE,
            supports=(ALI_TUB_SHOWER, ALI_DOUBLE_VANITY, ARNPRIOR_BATH),
            hero_on_left=True,
        ),
        add_editorial_page(
            filename="page-04-bathrooms-b.png",
            section="Bathrooms",
            title="Secondary-suite bathrooms still deserve the same complete, modern finish language.",
            body="Carleton, Bathroom Renovation, and Krishna project photos broaden the bathroom section without relying on the same repeated project set.",
            hero=CARLETON_NEW_BATH_B,
            supports=(CARLETON_NEW_BATH_A, BATHREN_VANITY, KRISHNA_BATH),
            hero_on_left=False,
        ),
        add_editorial_page(
            filename="page-05-kitchens.png",
            section="Kitchens & Kitchenettes",
            title="From full kitchens to compact suites, the finish quality stays bright, clean, and consistent.",
            body="Cabinet lines, durable counters, and easy-maintenance flooring help these kitchens feel practical, modern, and ready for everyday use.",
            hero=ASIYA_KITCHEN,
            supports=(ARNPRIOR_KITCHEN, CARLETON_NEW_KITCHENETTE_DETAIL, CARLETON_MORE_KITCHENETTE_DETAIL),
            hero_on_left=True,
        ),
        add_editorial_page(
            filename="page-06-basements-living.png",
            section="Basements",
            title="Lower-level spaces should feel bright, usable, and fully integrated with the rest of the home.",
            body="Finished bedrooms, living areas, and basement rooms add real variety to the portfolio while staying firmly in completed-space territory.",
            hero=ASIYA_LIVING,
            supports=(ASIYA_BEDROOM, ARBI_ROOM_A, ARBI_ROOM_B),
            hero_on_left=False,
        ),
        add_editorial_page(
            filename="page-07-lower-level-suites.png",
            section="Basement Suites",
            title="Lower-level suites should feel bright, efficient, and ready to live in from day one.",
            body="This spread focuses on practical suite planning: clean flooring, finished bath space, connected circulation, and a kitchenette that keeps the layout usable.",
            hero=CARLETON_NEW_KITCHENETTE_WIDE,
            supports=(CARLETON_NEW_BATH_A, CARLETON_MORE_HALL, CARLETON_MORE_KITCHENETTE_WIDE),
            hero_on_left=True,
        ),
        add_editorial_page(
            filename="page-08-flooring.png",
            section="Flooring",
            title="Flooring and stair work shown with more restraint so the craftsmanship reads first.",
            body="These images focus on landings, stair runs, surface continuity, and commercial flooring scale without repeating whole-room shots already used elsewhere.",
            hero=MISTY_STAIRS_DOWN,
            supports=(MISTY_STAIRS_UP, MISTY_STAIRS_LANDING, INDIAN_EMBASSY_FLOOR),
            hero_on_left=False,
        ),
        add_editorial_page(
            filename="page-09-final-details.png",
            section="Final Details",
            title="Practical rooms and support spaces help the portfolio feel real, complete, and professionally finished.",
            body="Fireplace walls, mudrooms, compact baths, and exterior detail work show the range of completed renovations without falling back on duplicate feature images.",
            hero=ASIYA_FIREPLACE,
            supports=(ALI_MUDROOM, BATHREN_COMPACT, ARNPRIOR_PORCH),
            hero_on_left=True,
        ),
        add_closing_page(),
    ]


def add_board_image(doc: Document, image_path: Path) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.32))


def build_document() -> Path:
    doc = Document()
    configure_document(doc)

    page_assets = build_page_assets()
    for index, asset in enumerate(page_assets):
        add_board_image(doc, asset)
        if index != len(page_assets) - 1:
            doc.add_page_break()

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def main() -> None:
    ensure_dirs()
    verify_inputs()
    output_path = build_document()
    print(output_path)


if __name__ == "__main__":
    main()
