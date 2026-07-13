from __future__ import annotations

import hashlib
import math
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT_DIR = Path(r"C:\Users\samra\Dropbox\CompanyCam Projects")
OUTPUT_DIR = ROOT_DIR / "Portfolio Documents"
ASSET_DIR = OUTPUT_DIR / "_final_assets"

OUTPUT_DOCX = OUTPUT_DIR / "final-projects-portfolio-by-category.docx"

COMPANY_NAME = "Ottawa Renovation Pro Ltd."
PORTFOLIO_TITLE = "Finished Project Portfolio"
PORTFOLIO_SUBTITLE = "Selected Renovation Work"
ISSUE_DATE = "July 2026"

PROJECTS = [
    "Bluegill",
    "Misty",
    "Arnprior",
    "Asiya",
    "Carleton",
    "Orleans",
    "Ali Bathroom",
    "Gimli",
    "Krishna Bathroom",
    "Penny",
]

MAX_IMAGES_PER_PROJECT = 2

DISPLAY_TITLES = {
    "Bluegill": "Basement Flooring Refresh",
    "Misty": "Entry Tile & Flooring",
    "Arnprior": "Kitchen & Bath Renovation",
    "Asiya": "Feature Wall & Living Area",
    "Carleton": "Basement Suite Update",
    "Orleans": "Orleans Powder Room",
    "Ali Bathroom": "Primary Bathroom Refresh",
    "Gimli": "Contemporary Shower Bath",
    "Krishna Bathroom": "Compact Modern Bath",
    "Penny": "Small Bath Update",
}

COVER_PROJECT_ORDER = [
    "Asiya",
    "Arnprior",
    "Ali Bathroom",
    "Gimli",
    "Misty",
    "Krishna Bathroom",
    "Orleans",
    "Penny",
]

CATEGORY_GROUPS = [
    (
        "Bathrooms & Powder Rooms",
        "Completed bathroom renovations, vanity updates, shower work, and compact washroom refreshes.",
        ["Orleans", "Ali Bathroom", "Gimli", "Krishna Bathroom", "Penny"],
    ),
    (
        "Kitchens & Multi-Room Updates",
        "Kitchen-focused and mixed-scope interior improvements with upgraded finishes and layout polish.",
        ["Arnprior"],
    ),
    (
        "Basement & Living Spaces",
        "Finished lower-level spaces, feature walls, and clean living-area transformations.",
        ["Bluegill", "Asiya", "Carleton"],
    ),
    (
        "Flooring & Entryways",
        "Entry, tile, and flooring projects that elevate circulation spaces and daily-use rooms.",
        ["Misty"],
    ),
]

PAGE_BG = "#F6F3EC"
CARD_BG = "#FFFFFF"
CARD_BORDER = "#D9D2C7"
INK = "#183640"
MUTED = "#676D72"
ACCENT = "#B98538"
CHIP_BG = "#EFE5D2"
CHIP_BORDER = "#DDC8A3"
PHOTO_MAT = "#FBFAF7"

BOARD_SIZE = (1800, 920)
COVER_COLLAGE_SIZE = (1800, 980)
OUTER_MARGIN = 76
CARD_GAP = 42
TOP_BAND = 222
CARD_RADIUS = 28
CARD_PADDING = 24
CARD_SHADOW = "#EAE4D8"

FONT_REGULAR = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")

FILENAME_RE = re.compile(
    r"^(?P<index>\d+)\s+(?P<project>.+?)\s+(?P<view>\d+)\s+(?P<label>a_before|before|after|aftr)\s*$",
    re.IGNORECASE,
)

PROJECT_LABELS = {
    "bluegill": "Bluegill",
    "misty": "Misty",
    "arnprior": "Arnprior",
    "asiya": "Asiya",
    "carleton": "Carleton",
    "orleans": "Orleans",
    "ali bathroom": "Ali Bathroom",
    "gimli": "Gimli",
    "krishna bathroom": "Krishna Bathroom",
    "penny": "Penny",
}

PROJECT_SOURCE_FOLDERS = {
    "Bluegill": [
        r"Construction new\Before After Pics",
        "flooring bluegill",
        "Lighting BLUEGILL 2",
        r"Construction new\Flooring Bluegill",
        r"Construction new\Lighting BLUEGILL 2",
    ],
    "Misty": [
        r"Construction new\Before After Pics",
        "Flooring Misty f",
        r"Construction new\Flooring Misty falls",
    ],
    "Arnprior": [
        r"Construction new\Before After Pics",
        "Arnprior - Arnprior ON",
        "Bathroom flooring kitchen Armprior",
        r"Construction new\Arnprior - Arnprior ON",
        r"Construction new\Bathroom flooring kitchen Armprior",
    ],
    "Asiya": [
        r"Construction new\Before After Pics",
        "Basement Kitchen Bathroom Asiya new",
        r"Construction new\Basement Kitchen Bathroom Asiya new",
    ],
    "Carleton": [
        r"Construction new\Before After Pics",
        "11 Carleton - Carleton Place Ontario",
        "11 Carleton More",
        "Basement Carleton New",
        "basement careltion",
        "Dhaval - Carleton Place Ontario",
        "Carleton Place 2 - Carleton Place Ontario",
        r"Construction new\11 Carleton",
        r"Construction new\11 Carleton More",
        r"Construction new\Basement Carleton New",
    ],
    "Orleans": [
        r"Construction new\Before After Pics",
        "09 Orleans - Ottawa Ontario",
        "Orleans - Ottawa Ontario",
        "26 orleans kitchen",
        r"Construction new\09 Orleans",
        r"Construction new\Arun - 14 N Harrow Dr Ottawa ON",
    ],
    "Ali Bathroom": [
        r"Construction new\Before After Pics",
        "Ali - Orleans Ottawa ON",
    ],
    "Gimli": [
        r"Construction new\Before After Pics",
        "Gimli - 9 Gimli Ct Ottawa ON",
        "Bathroom Gimli",
        r"Construction new\Bathroom Gimli",
    ],
    "Krishna Bathroom": [
        r"Construction new\Before After Pics",
        "Bathroom Krishna",
        r"Construction new\Bathroom Krishna",
    ],
    "Penny": [
        r"Construction new\Before After Pics",
        "Penny - 2960 Penny Dr Ottawa ON",
        "Flooring Penny",
        r"Construction new\Penny - 2960 Penny Dr Ottawa ON",
        r"Construction new\Flooring Bathroom Penny",
    ],
}

PROJECT_MATCH_TOKENS = {
    "Bluegill": ["bluegill"],
    "Misty": ["misty"],
    "Arnprior": ["arnprior"],
    "Asiya": ["asiya"],
    "Carleton": ["carleton"],
    "Orleans": ["orleans"],
    "Ali Bathroom": ["ali bathroom"],
    "Gimli": ["gimli"],
    "Krishna Bathroom": ["krishna"],
    "Penny": ["penny"],
}

CUSTOM_FINAL_IMAGES = {
    "Asiya": [
        ROOT_DIR / r"Construction new\Basement Kitchen Bathroom Asiya new\1a2.jpg",
    ],
    "Penny": [
        ROOT_DIR / r"Penny - 2960 Penny Dr Ottawa ON\project_99478548_07012025_0850.jpg",
    ],
}

PROJECT_CROP_OVERRIDES = {
    ("Orleans", "after"): (0.0, 0.06, 0.0, 0.12),
}


@dataclass
class Pair:
    project: str
    view_number: int
    before_path: Path
    after_path: Path


def load_font(path: Path, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    try:
        return ImageFont.truetype(str(path), size=size)
    except OSError:
        return ImageFont.load_default()


FONT_SMALL = load_font(FONT_BOLD, 22)
FONT_BODY = load_font(FONT_REGULAR, 28)
FONT_LABEL = load_font(FONT_BOLD, 44)
FONT_TITLE = load_font(FONT_BOLD, 56)


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def normalize_label(raw_label: str) -> str:
    label = raw_label.strip().lower()
    if label in {"a_before", "before"}:
        return "before"
    if label in {"after", "aftr"}:
        return "after"
    raise ValueError(f"Unknown label: {raw_label}")


def matches_project(project_name: str, raw_project: str) -> bool:
    normalized = raw_project.strip().lower()
    return any(token in normalized for token in PROJECT_MATCH_TOKENS[project_name])


def file_sha1(path: Path) -> str:
    hasher = hashlib.sha1()
    with path.open("rb") as source:
        while chunk := source.read(1024 * 1024):
            hasher.update(chunk)
    return hasher.hexdigest()


def parse_pairs() -> dict[str, list[Pair]]:
    projects: dict[str, list[Pair]] = {}

    for project_name, source_folders in PROJECT_SOURCE_FOLDERS.items():
        grouped: dict[int, dict[str, Path]] = {}
        seen_hashes: set[str] = set()
        for folder_name in source_folders:
            folder = ROOT_DIR / folder_name
            if not folder.exists():
                continue
            for path in sorted(folder.rglob("*")):
                if not path.is_file() or path.suffix.lower() not in {".jpg", ".jpeg", ".png", ".webp"}:
                    continue
                stem = re.sub(r"\s+", " ", path.stem.strip())
                match = FILENAME_RE.match(stem)
                if not match or not matches_project(project_name, match.group("project")):
                    continue

                file_hash = file_sha1(path)
                if file_hash in seen_hashes:
                    continue
                seen_hashes.add(file_hash)

                view_number = int(match.group("view"))
                label = normalize_label(match.group("label"))
                grouped.setdefault(view_number, {})[label] = path

        project_pairs = []
        for view_number, paths in grouped.items():
            before_path = paths.get("before")
            after_path = paths.get("after")
            if before_path and after_path:
                project_pairs.append(
                    Pair(
                        project=project_name,
                        view_number=view_number,
                        before_path=before_path,
                        after_path=after_path,
                    )
                )
        project_pairs.sort(key=lambda pair: pair.view_number)
        if project_pairs:
            projects[project_name] = project_pairs

    return projects


def select_pairs(project_pairs: list[Pair], max_pairs: int = MAX_IMAGES_PER_PROJECT) -> list[Pair]:
    if len(project_pairs) <= max_pairs:
        return project_pairs
    positions = [round(i * (len(project_pairs) - 1) / (max_pairs - 1)) for i in range(max_pairs)]
    ordered_unique = []
    seen = set()
    for pos in positions:
        if pos not in seen:
            seen.add(pos)
            ordered_unique.append(project_pairs[pos])
    return ordered_unique


def collect_final_images() -> list[tuple[str, list[Path]]]:
    pairs = parse_pairs()
    selected: list[tuple[str, list[Path]]] = []
    for project_name in PROJECTS:
        custom_images = [path for path in CUSTOM_FINAL_IMAGES.get(project_name, []) if path.exists()]
        if custom_images:
            selected.append((project_name, custom_images[:MAX_IMAGES_PER_PROJECT]))
            continue

        project_pairs = pairs.get(project_name, [])
        if not project_pairs:
            raise SystemExit(f"Missing after images for {project_name}")

        images = [pair.after_path for pair in select_pairs(project_pairs)]
        deduped: list[Path] = []
        seen_hashes: set[str] = set()
        for image_path in images:
            file_hash = file_sha1(image_path)
            if file_hash in seen_hashes:
                continue
            seen_hashes.add(file_hash)
            deduped.append(image_path)
        if not deduped:
            raise SystemExit(f"No usable after images found for {project_name}")
        selected.append((project_name, deduped[:MAX_IMAGES_PER_PROJECT]))
    return selected


def _set_rfonts(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def set_run_style(run, *, font_name: str, size: int, color: str, bold: bool = False, italic: bool = False) -> None:
    _set_rfonts(run, font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 0, line: float = 1.0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number_field(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    _set_rfonts(run, "Arial")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(placeholder)
    r_element.append(fld_end)


def set_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)

    label_run = paragraph.add_run(f"{COMPANY_NAME} | Page ")
    set_run_style(label_run, font_name="Arial", size=9, color=MUTED)
    add_page_number_field(paragraph)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("232B30")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def add_cover_page(doc: Document, cover_collage_path: Path, project_count: int) -> None:
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=44, line=1.0)

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(company, before=0, after=8, line=1.0)
    company_run = company.add_run(COMPANY_NAME.upper())
    set_run_style(company_run, font_name="Arial", size=11, color=ACCENT, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=0, after=6, line=1.0)
    title_run = title.add_run(PORTFOLIO_TITLE)
    set_run_style(title_run, font_name="Arial", size=25, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, before=0, after=16, line=1.15)
    subtitle_run = subtitle.add_run(PORTFOLIO_SUBTITLE)
    set_run_style(subtitle_run, font_name="Arial", size=13, color=MUTED)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(summary, before=0, after=20, line=1.2)
    summary_run = summary.add_run(
        f"A curated visual portfolio of {project_count} completed renovation projects organized by work type."
    )
    set_run_style(summary_run, font_name="Arial", size=11, color="#49535A")

    collage = doc.add_paragraph()
    collage.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(collage, before=0, after=18, line=1.0)
    collage.add_run().add_picture(str(cover_collage_path), width=Inches(6.15))

    issue = doc.add_paragraph()
    issue.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(issue, before=0, after=0, line=1.0)
    issue_run = issue.add_run(ISSUE_DATE)
    set_run_style(issue_run, font_name="Arial", size=10, color=MUTED, italic=True)

    doc.add_page_break()


def rounded_rect(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str) -> None:
    draw.rounded_rectangle(box, radius=CARD_RADIUS, fill=fill, outline=outline, width=3)


def draw_chip(draw: ImageDraw.ImageDraw, label: str, x: int, y: int) -> None:
    text_width = math.ceil(draw.textlength(label, font=FONT_SMALL))
    chip_box = (x, y, x + text_width + 34, y + 44)
    draw.rounded_rectangle(chip_box, radius=20, fill=CHIP_BG, outline=CHIP_BORDER, width=2)
    draw.text((x + 17, y + 10), label, font=FONT_SMALL, fill=ACCENT)


def crop_box_for(project_name: str, label: str) -> tuple[float, float, float, float] | None:
    return PROJECT_CROP_OVERRIDES.get((project_name, label.lower()))


def fit_image(image_path: Path, target_box: tuple[int, int], *, project_name: str, label: str) -> Image.Image:
    with Image.open(image_path) as src:
        image = ImageOps.exif_transpose(src).convert("RGB")

    crop = crop_box_for(project_name, label)
    if crop:
        left, top, right, bottom = crop
        width, height = image.size
        image = image.crop(
            (
                int(width * left),
                int(height * top),
                int(width * (1 - right)),
                int(height * (1 - bottom)),
            )
        )

    fitted = ImageOps.contain(image, target_box)
    canvas = Image.new("RGB", target_box, PHOTO_MAT)
    x = (target_box[0] - fitted.width) // 2
    y = (target_box[1] - fitted.height) // 2
    canvas.paste(fitted, (x, y))
    return canvas


def make_project_board(project_name: str, image_paths: list[Path], project_number: int) -> Path:
    board_path = ASSET_DIR / f"{project_number:02d}-{project_name.lower().replace(' ', '-')}-board.png"
    canvas = Image.new("RGB", BOARD_SIZE, PAGE_BG)
    draw = ImageDraw.Draw(canvas)

    display_name = DISPLAY_TITLES.get(project_name, project_name)

    draw_chip(draw, f"PROJECT {project_number:02d}", OUTER_MARGIN, 52)
    title_bbox = draw.textbbox((0, 0), display_name, font=FONT_TITLE)
    title_y = 118 - title_bbox[1]
    draw.text((OUTER_MARGIN, title_y), display_name, font=FONT_TITLE, fill=INK)

    subtitle_bbox = draw.textbbox((0, 0), "Selected final views", font=FONT_BODY)
    subtitle_y = title_y + (title_bbox[3] - title_bbox[1]) + 12 - subtitle_bbox[1]
    draw.text((OUTER_MARGIN, subtitle_y), "Selected final views", font=FONT_BODY, fill=MUTED)

    image_top = TOP_BAND
    card_height = BOARD_SIZE[1] - image_top - OUTER_MARGIN

    if len(image_paths) == 1:
        with Image.open(image_paths[0]) as src:
            portrait = src.height > src.width * 1.15
        if portrait:
            card_width = 1020
            left = (BOARD_SIZE[0] - card_width) // 2
            card_box = (
                left,
                image_top,
                left + card_width,
                image_top + card_height,
            )
        else:
            card_box = (
                OUTER_MARGIN,
                image_top,
                BOARD_SIZE[0] - OUTER_MARGIN,
                image_top + card_height,
            )
        shadow_box = (card_box[0] + 10, card_box[1] + 10, card_box[2] + 10, card_box[3] + 10)
        rounded_rect(draw, shadow_box, CARD_SHADOW, CARD_SHADOW)
        rounded_rect(draw, card_box, CARD_BG, CARD_BORDER)
        inner_size = (
            card_box[2] - card_box[0] - (CARD_PADDING * 2),
            card_box[3] - card_box[1] - (CARD_PADDING * 2),
        )
        fitted = fit_image(image_paths[0], inner_size, project_name=project_name, label="after")
        canvas.paste(fitted, (card_box[0] + CARD_PADDING, card_box[1] + CARD_PADDING))
    else:
        card_width = (BOARD_SIZE[0] - (OUTER_MARGIN * 2) - CARD_GAP) // 2
        for idx, image_path in enumerate(image_paths[:2]):
            left = OUTER_MARGIN + idx * (card_width + CARD_GAP)
            card_box = (left, image_top, left + card_width, image_top + card_height)
            shadow_box = (card_box[0] + 10, card_box[1] + 10, card_box[2] + 10, card_box[3] + 10)
            rounded_rect(draw, shadow_box, CARD_SHADOW, CARD_SHADOW)
            rounded_rect(draw, card_box, CARD_BG, CARD_BORDER)
            inner_size = (
                card_box[2] - card_box[0] - (CARD_PADDING * 2),
                card_box[3] - card_box[1] - (CARD_PADDING * 2),
            )
            fitted = fit_image(image_path, inner_size, project_name=project_name, label="after")
            canvas.paste(fitted, (card_box[0] + CARD_PADDING, card_box[1] + CARD_PADDING))

    canvas.save(board_path, format="PNG")
    return board_path


def make_cover_collage(projects: list[tuple[str, list[Path]]]) -> Path:
    collage_path = ASSET_DIR / "cover-collage.png"
    canvas = Image.new("RGB", COVER_COLLAGE_SIZE, PAGE_BG)
    draw = ImageDraw.Draw(canvas)

    columns = 4
    rows = 2
    tile_gap = 24
    inner_margin = 48
    tile_width = (COVER_COLLAGE_SIZE[0] - (inner_margin * 2) - (tile_gap * (columns - 1))) // columns
    tile_height = (COVER_COLLAGE_SIZE[1] - (inner_margin * 2) - (tile_gap * (rows - 1))) // rows

    project_map = dict(projects)
    images = []
    for project_name in COVER_PROJECT_ORDER:
        image_paths = project_map.get(project_name)
        if not image_paths:
            continue
        image_index = 1 if len(image_paths) > 1 and project_name == "Arnprior" else 0
        images.append(image_paths[image_index])
        if len(images) == columns * rows:
            break

    for idx, image_path in enumerate(images):
        row = idx // columns
        col = idx % columns
        left = inner_margin + col * (tile_width + tile_gap)
        top = inner_margin + row * (tile_height + tile_gap)
        tile_box = (left, top, left + tile_width, top + tile_height)
        rounded_rect(draw, tile_box, CARD_BG, CARD_BORDER)
        fitted = fit_image(image_path, (tile_width - 22, tile_height - 22), project_name="Cover", label="after")
        canvas.paste(fitted, (left + 11, top + 11))

    canvas.save(collage_path, format="PNG")
    return collage_path


def add_board_image(doc: Document, board_path: Path) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=0, after=12, line=1.0)
    paragraph.add_run().add_picture(str(board_path), width=Inches(6.18))


def add_category_heading(doc: Document, title_text: str, subtitle_text: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=0, after=4, line=1.0)
    title_run = title.add_run(title_text)
    set_run_style(title_run, font_name="Arial", size=20, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(subtitle, before=0, after=14, line=1.15)
    subtitle_run = subtitle.add_run(subtitle_text)
    set_run_style(subtitle_run, font_name="Arial", size=10, color=MUTED)


def add_contact_page(doc: Document) -> None:
    doc.add_page_break()

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=70, line=1.0)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, before=0, after=8, line=1.0)
    kicker_run = kicker.add_run(COMPANY_NAME.upper())
    set_run_style(kicker_run, font_name="Arial", size=11, color=ACCENT, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=0, after=12, line=1.0)
    title_run = title.add_run("Let's Build Your Next Space")
    set_run_style(title_run, font_name="Arial", size=24, color=INK, bold=True)

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(body, before=0, after=24, line=1.2)
    body_run = body.add_run(
        "Renovation planning, bathrooms, kitchens, basement finishing, and flooring updates for Ottawa-area homes."
    )
    set_run_style(body_run, font_name="Arial", size=11, color=MUTED)

    for line in ["Phone: +1 (613) 413-1756", "Website: www.ottawa-renovation.ca"]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=0, after=8, line=1.0)
        run = paragraph.add_run(line)
        set_run_style(run, font_name="Arial", size=13, color=INK, bold=True)


def build_document(projects: list[tuple[str, list[Path]]]) -> Path:
    doc = Document()
    configure_document(doc)

    cover_collage_path = make_cover_collage(projects)
    add_cover_page(doc, cover_collage_path, len(projects))

    project_map = dict(projects)
    project_number = 1
    first_category = True
    boards_on_page = 0
    for category_title, category_subtitle, category_projects in CATEGORY_GROUPS:
        if first_category:
            first_category = False
        elif boards_on_page == 2:
            doc.add_page_break()
            boards_on_page = 0
        else:
            spacer = doc.add_paragraph()
            set_paragraph_spacing(spacer, before=0, after=8, line=1.0)

        add_category_heading(doc, category_title, category_subtitle)
        for project_name in category_projects:
            image_paths = project_map[project_name]
            board_path = make_project_board(project_name, image_paths, project_number)
            add_board_image(doc, board_path)
            project_number += 1
            boards_on_page += 1
            if boards_on_page == 2 and project_name != category_projects[-1]:
                doc.add_page_break()
                boards_on_page = 0

    add_contact_page(doc)

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def main() -> None:
    ensure_dirs()
    projects = collect_final_images()
    output_path = build_document(projects)
    print(output_path)


if __name__ == "__main__":
    main()
