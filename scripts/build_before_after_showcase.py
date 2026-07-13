from __future__ import annotations

import math
import re
import hashlib
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT_DIR = Path(r"C:\Users\samra\Dropbox\CompanyCam Projects")
OUTPUT_DIR = ROOT_DIR / "Portfolio Documents"
ASSET_DIR = OUTPUT_DIR / "_assets"

SERIES_TITLE = "Modern Renovation Showcase"
SERIES_SUBTITLE = "Before & After Project Portfolio"
ISSUE_DATE = "July 2026"

VOLUMES = [
    ("Volume 01", ["Bluegill", "Misty"]),
    ("Volume 02", ["Arnprior", "Asiya"]),
    ("Volume 03", ["Carleton", "Orleans"]),
    ("Volume 04", ["Ali Bathroom", "Gimli"]),
    ("Volume 05", ["Krishna Bathroom", "Penny"]),
]

MAX_PAIRS_PER_PROJECT = 2

PAGE_BG = "#F6F4EE"
CARD_BG = "#FFFFFF"
CARD_BORDER = "#D7D3CC"
INK = "#183642"
MUTED = "#6A7075"
ACCENT = "#BE8B3D"
CHIP_BG = "#E6EEF0"
CHIP_BORDER = "#B8C9CE"

CANVAS_SIZE = (1800, 1120)
OUTER_MARGIN = 70
CARD_GAP = 48
CARD_LABEL_HEIGHT = 74
CARD_PADDING = 32

FONT_REGULAR = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")

FILENAME_RE = re.compile(
    r"^(?P<index>\d+)\s+(?P<project>.+?)\s+(?P<view>\d+)\s+(?P<label>a_before|before|after|aftr)\s*$",
    re.IGNORECASE,
)

PROJECT_LABELS = {
    "bluegill": "Bluegill",
    "misty": "Misty",
    "arnprior": "Arnprior",
    "asiya": "Asiya",
    "carleton": "Carleton",
    "orleans": "Orleans",
    "ali bathroom": "Ali Bathroom",
    "gimli": "Gimli",
    "krishna bathroom": "Krishna Bathroom",
    "penny": "Penny",
}

PROJECT_SOURCE_FOLDERS = {
    "Bluegill": [
        r"Construction new\Before After Pics",
        "flooring bluegill",
        "Lighting BLUEGILL 2",
        r"Construction new\Flooring Bluegill",
        r"Construction new\Lighting BLUEGILL 2",
    ],
    "Misty": [
        r"Construction new\Before After Pics",
        "Flooring Misty f",
        r"Construction new\Flooring Misty falls",
    ],
    "Arnprior": [
        r"Construction new\Before After Pics",
        "Arnprior - Arnprior ON",
        "Bathroom flooring kitchen Armprior",
        r"Construction new\Arnprior - Arnprior ON",
        r"Construction new\Bathroom flooring kitchen Armprior",
    ],
    "Asiya": [
        r"Construction new\Before After Pics",
        "Basement Kitchen Bathroom Asiya new",
        r"Construction new\Basement Kitchen Bathroom Asiya new",
    ],
    "Carleton": [
        r"Construction new\Before After Pics",
        "11 Carleton - Carleton Place Ontario",
        "11 Carleton More",
        "Basement Carleton New",
        "basement careltion",
        "Dhaval - Carleton Place Ontario",
        "Carleton Place 2 - Carleton Place Ontario",
        r"Construction new\11 Carleton",
        r"Construction new\11 Carleton More",
        r"Construction new\Basement Carleton New",
    ],
    "Orleans": [
        r"Construction new\Before After Pics",
        "09 Orleans - Ottawa Ontario",
        "Orleans - Ottawa Ontario",
        "26 orleans kitchen",
        r"Construction new\09 Orleans",
        r"Construction new\Arun - 14 N Harrow Dr Ottawa ON",
    ],
    "Ali Bathroom": [
        r"Construction new\Before After Pics",
        "Ali - Orleans Ottawa ON",
    ],
    "Gimli": [
        r"Construction new\Before After Pics",
        "Gimli - 9 Gimli Ct Ottawa ON",
        "Bathroom Gimli",
        r"Construction new\Bathroom Gimli",
    ],
    "Krishna Bathroom": [
        r"Construction new\Before After Pics",
        "Bathroom Krishna",
        r"Construction new\Bathroom Krishna",
    ],
    "Penny": [
        r"Construction new\Before After Pics",
        "Penny - 2960 Penny Dr Ottawa ON",
        "Flooring Penny",
        r"Construction new\Penny - 2960 Penny Dr Ottawa ON",
        r"Construction new\Flooring Bathroom Penny",
    ],
}

PROJECT_MATCH_TOKENS = {
    "Bluegill": ["bluegill"],
    "Misty": ["misty"],
    "Arnprior": ["arnprior"],
    "Asiya": ["asiya"],
    "Carleton": ["carleton"],
    "Orleans": ["orleans"],
    "Ali Bathroom": ["ali bathroom"],
    "Gimli": ["gimli"],
    "Krishna Bathroom": ["krishna"],
    "Penny": ["penny"],
}

CUSTOM_PROJECT_PAIRS = {
    "Asiya": [
        (
            ROOT_DIR / r"Construction new\Basement Kitchen Bathroom Asiya new\1b1.jpg",
            ROOT_DIR / r"Construction new\Basement Kitchen Bathroom Asiya new\1a2.jpg",
        )
    ],
    "Penny": [
        (
            ROOT_DIR / r"Penny - 2960 Penny Dr Ottawa ON\project_99478548_02022025_1233.jpg",
            ROOT_DIR / r"Penny - 2960 Penny Dr Ottawa ON\project_99478548_07012025_0850.jpg",
        )
    ],
}

PROJECT_CROP_OVERRIDES = {
    ("Orleans", "before"): (0.0, 0.0, 0.0, 0.12),
    ("Orleans", "after"): (0.0, 0.0, 0.0, 0.12),
}


@dataclass
class Pair:
    project: str
    view_number: int
    before_path: Path
    after_path: Path


def load_font(path: Path, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    try:
        return ImageFont.truetype(str(path), size=size)
    except OSError:
        return ImageFont.load_default()


FONT_SMALL = load_font(FONT_BOLD, 28)
FONT_LABEL = load_font(FONT_BOLD, 46)


def normalize_label(raw_label: str) -> str:
    label = raw_label.strip().lower()
    if label in {"a_before", "before"}:
        return "before"
    if label in {"after", "aftr"}:
        return "after"
    raise ValueError(f"Unknown label: {raw_label}")


def normalize_project_name(raw_project: str) -> str:
    cleaned = re.sub(r"\s+", " ", raw_project.strip()).lower()
    return PROJECT_LABELS.get(cleaned, raw_project.strip().title())


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def matches_project(project_name: str, raw_project: str) -> bool:
    normalized = raw_project.strip().lower()
    return any(token in normalized for token in PROJECT_MATCH_TOKENS[project_name])


def file_sha1(path: Path) -> str:
    hasher = hashlib.sha1()
    with path.open("rb") as source:
        while chunk := source.read(1024 * 1024):
            hasher.update(chunk)
    return hasher.hexdigest()


def parse_pairs() -> dict[str, list[Pair]]:
    projects: dict[str, list[Pair]] = {}

    for project_name, source_folders in PROJECT_SOURCE_FOLDERS.items():
        grouped: dict[int, dict[str, Path]] = {}
        seen_hashes: set[str] = set()
        for folder_name in source_folders:
            folder = ROOT_DIR / folder_name
            if not folder.exists():
                continue
            for path in sorted(folder.rglob("*")):
                if not path.is_file() or path.suffix.lower() not in {".jpg", ".jpeg", ".png", ".webp", ".heic"}:
                    continue
                stem = re.sub(r"\s+", " ", path.stem.strip())
                match = FILENAME_RE.match(stem)
                if not match or not matches_project(project_name, match.group("project")):
                    continue

                file_hash = file_sha1(path)
                if file_hash in seen_hashes:
                    continue
                seen_hashes.add(file_hash)

                view_number = int(match.group("view"))
                label = normalize_label(match.group("label"))
                grouped.setdefault(view_number, {})[label] = path

        project_pairs = []
        for view_number, paths in grouped.items():
            before_path = paths.get("before")
            after_path = paths.get("after")
            if before_path and after_path:
                project_pairs.append(
                    Pair(
                        project=project_name,
                        view_number=view_number,
                        before_path=before_path,
                        after_path=after_path,
                    )
                )
        project_pairs.sort(key=lambda pair: pair.view_number)
        if project_pairs:
            projects[project_name] = project_pairs

    for project_name, pair_paths in CUSTOM_PROJECT_PAIRS.items():
        custom_pairs = []
        for idx, (before_path, after_path) in enumerate(pair_paths, start=1):
            if before_path.exists() and after_path.exists():
                custom_pairs.append(
                    Pair(
                        project=project_name,
                        view_number=idx,
                        before_path=before_path,
                        after_path=after_path,
                    )
                )
        if custom_pairs:
            projects[project_name] = custom_pairs

    return projects


def select_pairs(project_pairs: list[Pair], max_pairs: int = MAX_PAIRS_PER_PROJECT) -> list[Pair]:
    if len(project_pairs) <= max_pairs:
        return project_pairs
    positions = [round(i * (len(project_pairs) - 1) / (max_pairs - 1)) for i in range(max_pairs)]
    ordered_unique = []
    seen = set()
    for pos in positions:
        if pos not in seen:
            seen.add(pos)
            ordered_unique.append(project_pairs[pos])
    return ordered_unique


def _set_rfonts(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def set_run_style(run, *, font_name: str, size: int, color: str, bold: bool = False, italic: bool = False) -> None:
    _set_rfonts(run, font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 0, line: float = 1.0) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def configure_document(doc: Document, footer_label: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_footer(section, footer_label)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1F2933")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def set_footer(section, footer_label: str) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
    run = paragraph.add_run(footer_label + " | Page ")
    set_run_style(run, font_name="Arial", size=9, color=MUTED)
    add_page_number_field(paragraph)


def add_page_number_field(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    _set_rfonts(run, "Arial")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(placeholder)
    r_element.append(fld_end)


def add_cover_page(doc: Document, volume_title: str, projects: Iterable[str]) -> None:
    doc.add_paragraph()
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=42, line=1.0)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, before=0, after=8, line=1.0)
    kicker_run = kicker.add_run(SERIES_SUBTITLE.upper())
    set_run_style(kicker_run, font_name="Arial", size=11, color=ACCENT, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=0, after=6, line=1.0)
    title_run = title.add_run(SERIES_TITLE)
    set_run_style(title_run, font_name="Arial", size=24, color=INK, bold=True)

    volume = doc.add_paragraph()
    volume.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(volume, before=0, after=14, line=1.0)
    volume_run = volume.add_run(volume_title)
    set_run_style(volume_run, font_name="Arial", size=15, color=MUTED, bold=True)

    project_line = doc.add_paragraph()
    project_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(project_line, before=0, after=28, line=1.15)
    project_run = project_line.add_run("Featured projects: " + " • ".join(projects))
    set_run_style(project_run, font_name="Arial", size=12, color="4F5B62")

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(summary, before=0, after=4, line=1.25)
    summary_run = summary.add_run(
        "A curated set of before-and-after renovation visuals arranged in a clean modern presentation."
    )
    set_run_style(summary_run, font_name="Arial", size=11, color="4F5B62")

    issue = doc.add_paragraph()
    issue.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(issue, before=0, after=0, line=1.0)
    issue_run = issue.add_run(ISSUE_DATE)
    set_run_style(issue_run, font_name="Arial", size=10, color=MUTED, italic=True)

    doc.add_page_break()


def add_project_heading(doc: Document, project_name: str, continuation: bool = False) -> None:
    heading = doc.add_paragraph()
    set_paragraph_spacing(heading, before=0, after=3, line=1.0)
    heading.paragraph_format.keep_with_next = True
    title_run = heading.add_run(project_name)
    set_run_style(title_run, font_name="Arial", size=18, color=INK, bold=True)
    if continuation:
        note = heading.add_run(" (continued)")
        set_run_style(note, font_name="Arial", size=10, color=MUTED, italic=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(subtitle, before=0, after=8, line=1.0)
    subtitle.paragraph_format.keep_with_next = True
    subtitle_run = subtitle.add_run("Selected transformation views")
    set_run_style(subtitle_run, font_name="Arial", size=10, color=MUTED)


def add_board(doc: Document, board_path: Path, *, width_inches: float = 6.18, after_pt: float = 10) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=0, after=after_pt, line=1.0)
    paragraph.add_run().add_picture(str(board_path), width=Inches(width_inches))


def draw_card_background(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int]) -> None:
    draw.rounded_rectangle(box, radius=26, fill=CARD_BG, outline=CARD_BORDER, width=3)


def fit_image_on_card(
    image: Image.Image,
    target_box: tuple[int, int],
    *,
    project_name: str,
    label: str,
) -> Image.Image:
    source = ImageOps.exif_transpose(image).convert("RGB")
    crop = PROJECT_CROP_OVERRIDES.get((project_name, label.lower()))
    if crop:
        left, top, right, bottom = crop
        width, height = source.size
        source = source.crop(
            (
                int(width * left),
                int(height * top),
                int(width * (1 - right)),
                int(height * (1 - bottom)),
            )
        )
    fitted = ImageOps.contain(source, target_box)
    canvas = Image.new("RGB", target_box, "#FBFAF7")
    x = (target_box[0] - fitted.width) // 2
    y = (target_box[1] - fitted.height) // 2
    canvas.paste(fitted, (x, y))
    return canvas


def add_card(
    draw: ImageDraw.ImageDraw,
    canvas: Image.Image,
    *,
    project_name: str,
    label: str,
    image_path: Path,
    box: tuple[int, int, int, int],
) -> None:
    draw_card_background(draw, box)
    left, top, right, bottom = box
    label_width = draw.textlength(label, font=FONT_LABEL)
    draw.text(
        (left + (right - left - label_width) / 2, top + 18),
        label,
        font=FONT_LABEL,
        fill=INK,
    )

    inner_width = right - left - (CARD_PADDING * 2)
    inner_height = bottom - top - CARD_LABEL_HEIGHT - (CARD_PADDING * 2)
    with Image.open(image_path) as src:
        fitted = fit_image_on_card(src, (inner_width, inner_height), project_name=project_name, label=label)
    canvas.paste(fitted, (left + CARD_PADDING, top + CARD_LABEL_HEIGHT + CARD_PADDING))


def make_board(pair: Pair) -> Path:
    board_path = ASSET_DIR / f"{pair.project.lower().replace(' ', '-')}-view-{pair.view_number:02d}.png"
    canvas = Image.new("RGB", CANVAS_SIZE, PAGE_BG)
    draw = ImageDraw.Draw(canvas)

    card_width = (CANVAS_SIZE[0] - (OUTER_MARGIN * 2) - CARD_GAP) // 2
    card_top = 78
    card_height = CANVAS_SIZE[1] - card_top - OUTER_MARGIN
    left_box = (OUTER_MARGIN, card_top, OUTER_MARGIN + card_width, card_top + card_height)
    right_box = (
        OUTER_MARGIN + card_width + CARD_GAP,
        card_top,
        OUTER_MARGIN + (card_width * 2) + CARD_GAP,
        card_top + card_height,
    )

    add_card(draw, canvas, project_name=pair.project, label="BEFORE", image_path=pair.before_path, box=left_box)
    add_card(draw, canvas, project_name=pair.project, label="AFTER", image_path=pair.after_path, box=right_box)

    canvas.save(board_path, format="PNG")
    return board_path


def build_document(volume_title: str, project_names: list[str], available_pairs: dict[str, list[Pair]]) -> Path:
    doc = Document()
    configure_document(doc, volume_title)
    add_cover_page(doc, volume_title, project_names)

    boards_on_page = 0
    first_project = True
    for project_name in project_names:
        selected_pairs = select_pairs(available_pairs[project_name])
        board_paths = [make_board(pair) for pair in selected_pairs]
        project_starts_on_partial_page = boards_on_page == 1
        if not first_project and boards_on_page == 2:
            doc.add_page_break()
            boards_on_page = 0
            project_starts_on_partial_page = False
        add_project_heading(doc, project_name)
        first_project = False

        for idx, board_path in enumerate(board_paths):
            if idx == 0 and project_starts_on_partial_page:
                width_inches = 5.62
            else:
                width_inches = 6.18 if boards_on_page == 0 else 5.98
            add_board(doc, board_path, width_inches=width_inches)
            boards_on_page += 1
            if idx < len(board_paths) - 1 and boards_on_page == 2:
                doc.add_page_break()
                boards_on_page = 0
                add_project_heading(doc, project_name, continuation=True)
        if boards_on_page < 2:
            spacer = doc.add_paragraph()
            set_paragraph_spacing(spacer, before=0, after=2, line=1.0)

    filename = f"{volume_title.lower().replace(' ', '-')}-" + "-".join(
        project.lower().replace(" ", "-") for project in project_names
    )
    output_path = OUTPUT_DIR / f"{filename}.docx"
    doc.save(output_path)
    return output_path


def main() -> None:
    ensure_dirs()
    pairs = parse_pairs()

    missing = [
        project
        for _, volume_projects in VOLUMES
        for project in volume_projects
        if project not in pairs
    ]
    if missing:
        raise SystemExit(f"Missing complete before/after pairs for: {', '.join(missing)}")

    created: list[Path] = []
    for volume_title, project_names in VOLUMES:
        created.append(build_document(volume_title, project_names, pairs))

    print("Created documents:")
    for path in created:
        print(path)


if __name__ == "__main__":
    main()
